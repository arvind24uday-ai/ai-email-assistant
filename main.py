from googleapiclient.discovery import build
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
import os
import requests
import base64
from email.mime.text import MIMEText
from docx import Document
import datetime

# ✅ Gmail access (read + send)
SCOPES = ['https://www.googleapis.com/auth/gmail.modify']

AUTO_SEND = False  # 🔥 keep False unless you want auto replies


# 🔹 Extract name safely
def extract_name(from_field):
    if "<" in from_field:
        name = from_field.split("<")[0].strip()
    else:
        name = from_field

    bad_names = ["google", "no-reply", "noreply", "support", "team"]
    first_name = name.split()[0] if name else "there"

    if first_name.lower() in bad_names:
        return "there"

    return first_name


# 🔹 AI classifier (FIXED LOGIC)
def is_important_email(subject):
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "llama3",
                "prompt": f"""
You are an email classifier.

IMPORTANT:
- Job / recruiter / HR
- Interview / application
- Personal / work messages
- Action required

NOT IMPORTANT:
- Ads
- Promotions
- Newsletters
- Spam

Reply ONLY with:
IMPORTANT
or
NOT IMPORTANT

Email subject: {subject}
""",
                "stream": False
            }
        )

        result = response.json().get('response', '').strip().upper()
        print("🧠 AI Decision:", result)

        return result == "IMPORTANT"   # ✅ FIXED

    except Exception as e:
        print("⚠️ AI Error:", e)
        return False


# 🔹 AI reply generator (CLEANED)
def generate_reply(subject, name):
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "llama3",
                "prompt": f"""
You are writing a real email reply.

STRICT RULES:
- DO NOT add explanations
- DO NOT say "Here is a reply"
- Only output the email

FORMAT:
Dear {name},

(3-4 lines reply)

Best regards,
Arvind Uday M

Email context: {subject}
""",
                "stream": False
            }
        )

        reply = response.json().get('response', '').strip()

        # 🔥 CLEAN unwanted lines
        bad_phrases = ["here is", "email reply", "response:", "reply:", "subject:"]

        for phrase in bad_phrases:
            if phrase in reply.lower():
                reply = reply.split("\n", 1)[-1].strip()

        return reply

    except Exception as e:
        return f"Error generating reply: {e}"


# 🔹 Gmail authentication
def authenticate():
    creds = None

    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)

    if not creds or not creds.valid:
        flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
        creds = flow.run_local_server(port=0)

        with open('token.json', 'w') as token:
            token.write(creds.to_json())

    return creds


# 🔹 Send reply
def send_reply(service, original_msg, reply_text):
    headers = original_msg['payload']['headers']

    subject = ""
    from_email = ""

    for header in headers:
        if header['name'] == 'Subject':
            subject = header['value']
        if header['name'] == 'From':
            from_email = header['value']

    message = MIMEText(reply_text)
    message['to'] = from_email
    message['subject'] = f"Re: {subject}"

    raw = base64.urlsafe_b64encode(message.as_bytes()).decode()

    service.users().messages().send(
        userId='me',
        body={'raw': raw}
    ).execute()

    print("📤 Reply sent!")


# 🔹 MAIN FUNCTION
def get_emails():
    creds = authenticate()
    service = build('gmail', 'v1', credentials=creds)

    results = service.users().messages().list(userId='me', maxResults=10).execute()
    messages = results.get('messages', [])

    doc = Document()
    doc.add_heading('Important Emails Report', 0)

    seen = set()
    count = 0

    for msg in messages:
        full_msg = service.users().messages().get(userId='me', id=msg['id']).execute()

        headers = full_msg['payload']['headers']
        subject = ""
        from_field = ""

        for header in headers:
            if header['name'] == 'Subject':
                subject = header['value']
            if header['name'] == 'From':
                from_field = header['value']

        if not subject or subject in seen:
            continue

        seen.add(subject)

        print("\n📩 Subject:", subject)

        # 🔥 FILTER
        if not is_important_email(subject):
            print("⛔ Skipped (not important)")
            continue

        print("✅ IMPORTANT email")

        name = extract_name(from_field)

        reply = generate_reply(subject, name)

        print("🤖 Suggested Reply:\n", reply)

        # 📄 Save to Word
        doc.add_heading(f"Email {count + 1}", level=1)
        doc.add_paragraph(f"Subject: {subject}")
        doc.add_paragraph(f"From: {from_field}")
        doc.add_paragraph("Reply:")
        doc.add_paragraph(reply)
        doc.add_paragraph("-" * 50)

        count += 1

        # 📤 Optional send
        if AUTO_SEND:
            send_reply(service, full_msg, reply)

    # 💾 SAFE FILE SAVE (FIXED ERROR)
    if count > 0:
        filename = f"important_emails_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        print(f"\n📄 Word document saved: {filename}")
    else:
        print("\n⚠️ No important emails found")


# 🚀 RUN
get_emails()